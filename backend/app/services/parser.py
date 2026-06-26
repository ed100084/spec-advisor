"""文件解析服務 - 支援 PDF, DOCX, XLSX"""
from pathlib import Path

import PyPDF2
from docx import Document as DocxDocument
from openpyxl import load_workbook


def parse_pdf(file_path: str) -> str:
    text_parts = []
    with open(file_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def parse_docx(file_path: str) -> str:
    doc = DocxDocument(file_path)
    text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                text_parts.append(row_text)

    return "\n".join(text_parts)


def parse_xlsx(file_path: str) -> str:
    wb = load_workbook(file_path, read_only=True, data_only=True)
    text_parts = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        text_parts.append(f"=== {sheet_name} ===")
        for row in ws.iter_rows(values_only=True):
            row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
            if row_text.strip(" |"):
                text_parts.append(row_text)

    wb.close()
    return "\n".join(text_parts)


def parse_file(file_path: str) -> str:
    suffix = Path(file_path).suffix.lower()
    parsers = {
        ".pdf": parse_pdf,
        ".docx": parse_docx,
        ".xlsx": parse_xlsx,
        ".xls": parse_xlsx,
    }
    parser = parsers.get(suffix)
    if not parser:
        raise ValueError(f"不支援的檔案格式: {suffix}")
    return parser(file_path)
